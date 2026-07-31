# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
import sqlite3
import stat
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = ROOT / "文档"
TARGET = DOC_ROOT / "全栈1班-张三-20230001-暑期实训材料"
PROJECT = "RAG职业技能培训系统"
TODAY = datetime.now().strftime("%Y年%m月%d日")
STAMP = datetime.now().strftime("%Y%m%d-%H%M%S")
BACKUP = DOC_ROOT / f"备份-全栈1班-张三-20230001-暑期实训材料-{STAMP}"

OUT_ACHIEVEMENTS = TARGET / "02-项目成果"
OUT_REPORTS = TARGET / "03-实训报告与归档材料"

COMMON_INFO = [
    ("姓名", "待填写"),
    ("学号", "待填写"),
    ("班级", "待填写"),
    ("实训方向/班级", "全栈开发 / 智能软件开发方向（待确认）"),
    ("项目名称", PROJECT),
    ("指导教师", "待填写"),
    ("Git仓库地址", "待填写"),
    ("提交日期", TODAY),
]

MODULES = [
    ("学习平台", "按岗位和课程展示知识内容，支持章节浏览与跳转练习。", "部分完成", "课程内容来自本地知识库与上传资料，仍需继续优化页面文案和知识点展示体验。"),
    ("智能问答", "基于课程知识库检索与 DeepSeek 大模型生成回答，切换功能时保持生成状态。", "部分完成", "依赖外部模型服务和网络代理；模型不可用时会影响回答质量。"),
    ("实操指导", "针对配置、安装、部署、排错等问题生成分步指导和 AI 示意图。", "部分完成", "已实现模板加模型生成和前端展示；示意图为简化图，不是真实截图。"),
    ("模拟练习与错题本", "从题库生成练习，提交后记录正确率与错题。", "部分完成", "题库 PDF/知识库上传已接入，仍需更多真实题库验证。"),
    ("社区问答", "用户可发布问题、回答、点赞、删除本人回复。", "部分完成", "核心交互已修复，仍需继续做权限和异常场景测试。"),
    ("个人中心与反馈", "展示账号信息、成绩变化，支持反馈提交和账号注销。", "部分完成", "成绩仅使用真实练习记录；无记录时不造假。"),
    ("后台管理", "提供用户、岗位、知识库、题库、社区、反馈和成绩管理。", "部分完成", "已按正常/封禁/注销分类管理用户；后台局部中文源码/界面仍存在乱码待修复。"),
]

TECH_ROWS = [
    ("前端技术", "React、TypeScript、Vite、Tailwind CSS、lucide-react", "React 19.2.7，Vite 8.1.1", "适合快速构建单页应用，组件化维护学习端和管理端界面。"),
    ("后端技术", "FastAPI、Uvicorn、Python", "FastAPI 0.115.0，Uvicorn 0.30.0", "接口开发效率高，类型清晰，便于与前端和大模型服务对接。"),
    ("数据库", "SQLite、SQLAlchemy", "本地 training.db.runtime", "适合本地实训部署和快速验证；SQLAlchemy 便于后续切换数据库。"),
    ("知识检索", "ChromaDB、PyMuPDF", "ChromaDB 0.5.5，PyMuPDF 1.24.0", "用于 PDF/知识库切块、存储和相似度检索。"),
    ("大模型服务", "DeepSeek API、httpx", "deepseek-chat，text-embedding-v2", "用于智能问答、实操步骤生成和部分内容理解。"),
    ("账号与通知", "JWT、SMTP 邮件验证码", "python-jose、aiosmtplib", "支持登录鉴权、注册验证码和忘记密码验证码。"),
    ("测试工具", "oxlint、TypeScript build、Python compileall", "本轮 pytest 未执行：当前环境缺少 pytest 模块", "用于静态检查、构建检查和后端语法检查。"),
]


def get_stats() -> dict[str, int | None]:
    db_path = ROOT / "backend" / "training.db.runtime"
    tables = [
        "users",
        "positions",
        "courses",
        "knowledge_points",
        "course_materials",
        "doc_chunks",
        "questions",
        "community_questions",
        "community_answers",
        "feedbacks",
        "qa_sessions",
        "qa_messages",
    ]
    stats: dict[str, int | None] = {}
    if not db_path.exists():
        return {table: None for table in tables}
    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()
        for table_name in tables:
            try:
                cur.execute(f"select count(*) from {table_name}")
                stats[table_name] = int(cur.fetchone()[0])
            except Exception:
                stats[table_name] = None
    finally:
        conn.close()
    return stats


def style_doc(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "黑体"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def para(doc: Document, text: str = ""):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    return p


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def set_cell_text(cell, text) -> None:
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str = "F3F4F6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[index], header)
        shade_cell(tbl.rows[0].cells[index])
    for row in rows:
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)


def save_doc(doc: Document, path: Path) -> None:
    style_doc(doc)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def info_section(doc: Document, extra: list[tuple[str, str]] | None = None) -> None:
    doc.add_heading("一、基本信息", level=1)
    rows = list(COMMON_INFO)
    if extra:
        rows.extend(extra)
    table(doc, ["项目", "内容"], rows)


def build_prd() -> None:
    doc = Document()
    title(doc, f"{PROJECT} PRD文档")
    doc.add_heading("一、文档基础信息", level=1)
    table(doc, ["项目", "内容"], [
        ("项目名称", PROJECT),
        ("文档版本", "V1.0"),
        ("产品定位", "面向职业技能学习的 RAG 知识培训与练习系统"),
        ("目标用户", "学员、教师/管理员"),
        ("核心目标", "围绕岗位知识库提供学习、问答、实操指导、练习和管理能力"),
        ("文档目的", "说明当前系统产品需求、完成情况和后续改进方向"),
        ("参考依据", "本地项目源码、数据库结构、答辩模板和实际验证结果"),
    ])
    doc.add_heading("二、产品核心价值", level=1)
    doc.add_heading("2.1 学员端价值", level=2)
    bullets(doc, [
        "按职业方向选择课程和知识库，降低学习入口复杂度。",
        "通过智能问答把知识库内容转化为可理解回答。",
        "通过模拟练习、错题本和成绩变化查看真实学习反馈。",
        "通过实操指导获得分步骤操作建议，适合程序员以外的职业技能问题。",
    ])
    doc.add_heading("2.2 管理端价值", level=2)
    bullets(doc, [
        "支持岗位、课程、知识库 PDF 和题库 PDF 的上传管理。",
        "支持用户封禁、解封、注销账号找回和分类管理。",
        "支持社区、反馈和学习成绩查看，便于教师了解系统使用情况。",
    ])
    doc.add_heading("三、核心功能模块划分", level=1)
    rows = [("学员端", item[0], item[1], item[2]) for item in MODULES[:6]]
    rows.append(("管理端", "后台管理", MODULES[6][1], MODULES[6][2]))
    table(doc, ["端", "模块", "核心功能概述", "完成情况"], rows)
    doc.add_heading("四、详细功能需求", level=1)
    for item in MODULES:
        doc.add_heading(item[0], level=2)
        para(doc, f"{item[1]}当前状态：{item[2]}。说明：{item[3]}")
    doc.add_heading("五、交互需求", level=1)
    bullets(doc, [
        "页面以登录页、岗位选择、学习空间、后台管理为主流程。",
        "长耗时智能问答生成过程中允许用户切换其它功能，回来后继续查看结果。",
        "上传、删除、封禁、注销等高风险操作应提供明确状态反馈。",
    ])
    doc.add_heading("六、非功能需求", level=1)
    table(doc, ["类别", "需求说明", "当前情况"], [
        ("性能", "本地启动和页面切换应尽量流畅", "首次启动仍偏慢，需后续优化依赖加载和服务启动流程。"),
        ("安全", "登录鉴权、密码重置、敏感配置不写入文档", "JWT 与验证码已接入；文档不包含真实密钥。"),
        ("兼容", "支持现代桌面浏览器", "主要在本地浏览器环境验证，移动端未作为重点。"),
        ("可扩展", "知识库和题库可按职业扩展", "已支持新增岗位后上传知识库和题库。"),
    ])
    doc.add_heading("七、迭代规划", level=1)
    table(doc, ["版本", "规划内容"], [
        ("V1.0", "完成学习、问答、练习、社区、反馈、后台管理等核心闭环。"),
        ("V1.1", "修复局部乱码、优化启动速度、完善异常提示和端到端测试。"),
        ("V1.2", "增强知识库解析质量、增加更多真实职业模板、补充部署和权限审计能力。"),
    ])
    doc.add_heading("八、验收标准", level=1)
    bullets(doc, [
        "能够本地启动前后端并访问学习端和管理端。",
        "能够上传知识库/题库并用于问答或练习。",
        "不使用伪造成绩或伪造运营数据，无数据时显示空状态。",
    ])
    save_doc(doc, OUT_ACHIEVEMENTS / f"{PROJECT}-PRD文档.docx")


def build_architecture() -> None:
    doc = Document()
    title(doc, "技术架构图和技术选型说明")
    info_section(doc)
    doc.add_heading("二、项目技术架构图", level=1)
    arch = (
        "用户浏览器\n"
        "  ↓ HTTP/Hash Router\n"
        "React + Vite 前端（学习端 / 管理端）\n"
        "  ↓ REST API + JWT\n"
        "FastAPI 后端路由层\n"
        "  ├─ Auth / Users / Positions / Learning / QA / Practical / Exercise / Community / Feedback / Admin\n"
        "  ├─ SQLAlchemy ORM → SQLite 本地数据库\n"
        "  ├─ PDF 导入与题库解析 → PyMuPDF / lxml\n"
        "  ├─ RAG 检索 → DocChunk + ChromaDB\n"
        "  ├─ DeepSeek API → 智能问答与实操指导生成\n"
        "  └─ SMTP → 注册与找回密码验证码"
    )
    run = doc.add_paragraph().add_run(arch)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_heading("三、技术选型说明", level=1)
    table(doc, ["类别", "选用技术/工具", "版本或说明", "选择原因"], TECH_ROWS)
    doc.add_heading("四、核心模块与技术对应关系", level=1)
    table(doc, ["序号", "核心模块", "主要技术", "说明"], [
        ("1", "智能问答", "FastAPI、ChromaDB、DeepSeek、React", "根据课程知识库检索上下文并生成回答。"),
        ("2", "知识库/题库上传", "PyMuPDF、SQLAlchemy、SQLite", "解析 PDF/文档内容并写入知识点、切片和题库。"),
        ("3", "模拟练习与成绩", "React、FastAPI、SQLite", "提交练习后记录真实正确率和错题。"),
        ("4", "后台管理", "React、JWT、FastAPI", "管理用户、岗位、资料、反馈和社区内容。"),
    ])
    doc.add_heading("五、学生确认", level=1)
    table(doc, ["确认项", "内容"], [("学生签名", "待填写"), ("日期", TODAY)])
    save_doc(doc, OUT_ACHIEVEMENTS / f"{PROJECT}-技术架构图和技术选型说明.docx")


def build_completion() -> None:
    doc = Document()
    title(doc, "毕业设计文字版完成说明")
    info_section(doc, [("毕业设计题目", PROJECT), ("联系电话", "待填写")])
    doc.add_heading("二、毕业设计总体完成情况", level=1)
    para(doc, f"{PROJECT} 当前已形成前后端分离的本地运行系统，包含学习端和管理端。系统围绕职业岗位知识库组织学习内容，接入 RAG 检索、DeepSeek 大模型、题库练习、社区问答和反馈管理等能力。整体完成度为“部分完成”：主流程已经可以用于答辩演示和本地验证，但仍存在局部乱码、启动速度偏慢、外部模型和邮箱服务依赖等待完善问题。")
    doc.add_heading("三、需求与功能完成情况", level=1)
    table(doc, ["序号", "功能模块", "完成情况", "说明"], [(str(i + 1), item[0], item[2], item[3]) for i, item in enumerate(MODULES)])
    doc.add_heading("四、存在问题与后续改进计划", level=1)
    table(doc, ["问题", "原因分析", "后续计划", "预计时间"], [
        ("局部中文乱码", "部分源码或文案在编辑过程中出现编码显示异常", "统一检查前后端中文文案和源码编码", "待填写"),
        ("启动时间偏慢", "前端开发服务器首次编译、后端加载数据库/向量库和代理检查均耗时", "优化启动脚本、增加已启动检测和生产构建运行方案", "待填写"),
        ("模型/邮箱服务依赖外部环境", "DeepSeek 和 SMTP 需要网络、代理和有效配置", "增加更清晰的配置检查、降级提示和本地 fallback", "待填写"),
        ("自动化测试不足", "当前测试环境缺少 pytest，前端缺少端到端测试脚本", "补齐测试依赖和关键流程自动化用例", "待填写"),
    ])
    doc.add_heading("五、学生确认", level=1)
    table(doc, ["学生签名", "日期"], [("待填写", TODAY)])
    save_doc(doc, OUT_ACHIEVEMENTS / f"{PROJECT}-毕业设计完成说明.docx")


def build_test_report(stats: dict[str, int | None]) -> None:
    doc = Document()
    title(doc, "项目测试报告")
    info_section(doc, [("测试人员", "待填写"), ("测试日期", TODAY)])
    doc.add_heading("二、测试目的", level=1)
    para(doc, "本次测试用于记录项目在文档生成时的真实验证情况，重点检查前端构建、静态检查、后端语法、服务端口连通、核心模块待测范围和已知风险，不伪造未执行的人工功能测试结果。")
    doc.add_heading("三、测试环境", level=1)
    table(doc, ["环境项", "说明"], [
        ("操作系统", "Windows 本地开发环境"),
        ("浏览器或客户端", "本地浏览器 / Codex 内置浏览器"),
        ("前端运行环境", "Node.js、npm、Vite 开发服务器，端口 5173"),
        ("后端运行环境", "Python、FastAPI、Uvicorn，端口 8000"),
        ("数据库", "SQLite 本地运行库 training.db.runtime，ChromaDB 本地向量库"),
        ("测试地址", "文档生成时 http://127.0.0.1:5173 与 http://127.0.0.1:8000/docs 未响应"),
    ])
    doc.add_heading("四、测试范围", level=1)
    table(doc, ["序号", "功能模块", "是否测试", "说明"], [
        ("1", "前端静态检查", "是", "npm.cmd run lint：通过"),
        ("2", "前端生产构建", "是", "npm.cmd run build：通过"),
        ("3", "后端语法检查", "是", "python -m compileall -q backend\\app：通过"),
        ("4", "后端 pytest", "否", "当前 Python 环境缺少 pytest 模块，未执行"),
        ("5", "端口连通", "是", "生成文档时前后端端口均无法连接，记录为未通过/服务未启动"),
        ("6", "核心 UI 流程", "待人工复测", "登录、问答、上传、练习、社区等需在服务启动后逐项复测"),
    ])
    doc.add_heading("五、功能测试用例", level=1)
    table(doc, ["用例编号", "功能模块", "测试步骤", "预期结果", "实际结果", "测试结果"], [
        ("TC-001", "前端静态检查", "执行 npm.cmd run lint", "无 lint 错误", "命令通过", "通过"),
        ("TC-002", "前端构建", "执行 npm.cmd run build", "生成 dist 构建产物", "命令通过，生成前端构建资源", "通过"),
        ("TC-003", "后端语法", "执行 python -m compileall -q backend\\app", "无语法错误", "命令通过", "通过"),
        ("TC-004", "后端 pytest", "执行 python -m pytest", "运行现有测试", "No module named pytest", "未通过/环境缺依赖"),
        ("TC-005", "服务连通", "请求 5173 与 8000/docs", "返回 HTTP 200", "生成文档时无法连接远程服务器", "未通过/服务未启动"),
        ("TC-006", "智能问答切换不中断", "服务启动后提问并切换功能板块", "返回后仍显示生成中或结果", "本轮未启动浏览器复测", "待复测"),
    ])
    doc.add_heading("六、接口或数据测试情况", level=1)
    table(doc, ["测试项", "测试方式", "测试结果", "说明"], [
        ("接口调用", "端口 HTTP 请求", "未通过", "文档生成时服务未运行或未响应。"),
        ("数据新增", "数据库统计与页面功能记录", "待复测", "已有本地数据，但本轮未新增验证。"),
        ("数据修改", "页面操作/接口调用", "待复测", "用户状态、反馈处理等需启动服务后验证。"),
        ("数据删除", "页面操作/接口调用", "待复测", "回答删除、错题删除等需启动服务后验证。"),
        ("权限校验", "不同角色登录测试", "待复测", "管理员不能封禁自己等逻辑已实现过，但本轮未重新执行。"),
    ])
    doc.add_heading("七、本地数据库统计", level=1)
    rows = [
        ("用户数", stats.get("users")),
        ("岗位数", stats.get("positions")),
        ("课程数", stats.get("courses")),
        ("知识点数", stats.get("knowledge_points")),
        ("资料数", stats.get("course_materials")),
        ("文档切片数", stats.get("doc_chunks")),
        ("题目数", stats.get("questions")),
        ("社区问题数", stats.get("community_questions")),
        ("社区回答数", stats.get("community_answers")),
        ("反馈数", stats.get("feedbacks")),
        ("问答会话数", stats.get("qa_sessions")),
        ("问答消息数", stats.get("qa_messages")),
    ]
    table(doc, ["统计项", "截至生成时本地数据库数量"], [(name, str(value)) for name, value in rows])
    save_doc(doc, OUT_ACHIEVEMENTS / f"{PROJECT}-项目测试报告.docx")


def build_deploy_doc() -> None:
    doc = Document()
    title(doc, "项目部署文档")
    info_section(doc, [("部署方式", "本地部署"), ("提交日期", TODAY)])
    doc.add_heading("二、项目目录说明", level=1)
    table(doc, ["目录或文件", "作用说明"], [
        ("frontend /", "React + Vite 前端源码、页面组件、API 封装和构建配置。"),
        ("backend /", "FastAPI 后端源码、路由、模型、服务、数据库和上传目录。"),
        ("backend/training.db.runtime", "本地 SQLite 运行库。"),
        ("backend/chroma_db /", "ChromaDB 本地向量数据库。"),
        ("知识库 /", "内置或上传的职业知识库资料。"),
        ("文档 /", "答辩模板、归档材料和生成后的项目文档。"),
        ("启动系统.bat", "一键启动前端和后端，并检查端口。"),
        ("停止系统.bat", "停止 8000 和 5173 端口对应进程。"),
    ])
    doc.add_heading("三、运行环境要求", level=1)
    table(doc, ["环境项", "要求或版本", "备注"], [
        ("操作系统", "Windows", "当前项目按 Windows 本地脚本编写。"),
        ("前端运行环境", "Node.js + npm", "package.json 中使用 Vite、React、TypeScript。"),
        ("后端运行环境", "Python", "依赖见 backend/requirements.txt。"),
        ("数据库", "SQLite + ChromaDB", "默认使用本地文件，无需单独安装数据库服务。"),
        ("端口占用", "5173、8000", "前端 Vite 端口 5173，后端 FastAPI 端口 8000。"),
        ("其他依赖", "DeepSeek API、SMTP、可选本地代理 127.0.0.1:7897", "用于智能问答、实操生成和验证码。"),
    ])
    doc.add_heading("四、配置文件说明", level=1)
    table(doc, ["配置文件路径", "配置项", "说明", "示例值"], [
        ("backend/.env", "DEEPSEEK_API_KEY", "DeepSeek 接口密钥，文档不写真实值", "******"),
        ("backend/.env", "DATABASE_TYPE", "数据库类型", "sqlite"),
        ("backend/.env", "JWT_SECRET_KEY", "JWT 签名密钥", "******"),
        ("backend/.env", "EMAIL_HOST / EMAIL_USER / EMAIL_PASSWORD", "邮箱验证码 SMTP 配置", "smtp.qq.com / ******"),
        ("frontend/.env 或默认配置", "VITE_API_BASE_URL", "前端接口地址", "http://localhost:8000/api"),
    ])
    doc.add_heading("五、数据库初始化步骤", level=1)
    bullets(doc, [
        "默认使用 backend/training.db.runtime；若文件存在，后端启动后直接连接。",
        "ChromaDB 数据位于 backend/chroma_db，知识库上传或导入后写入该目录。",
        "如果需要重新初始化，应先备份 training.db.runtime 和 chroma_db，再运行后端初始化逻辑或重新上传资料。",
    ])
    doc.add_heading("六、依赖安装步骤", level=1)
    table(doc, ["模块", "所在目录", "安装命令"], [
        ("前端", r"D:\毕设2\frontend", "npm install"),
        ("后端", r"D:\毕设2\backend", "pip install -r requirements.txt"),
        ("测试", r"D:\毕设2\backend", "如需 pytest，执行 pip install pytest 后再运行测试"),
    ])
    doc.add_heading("七、启动运行步骤", level=1)
    table(doc, ["步骤", "操作说明", "命令或截图说明"], [
        ("1", "确认 Python、Node.js、npm 可用", "python --version；npm --version"),
        ("2", "启动后端服务", r"cd /d D:\毕设2\backend && python -m uvicorn app.main:app --host 127.0.0.1 --port 8000"),
        ("3", "启动前端服务", r"cd /d D:\毕设2\frontend && npm.cmd run dev -- --host 127.0.0.1"),
        ("4", "或使用一键脚本", r"双击 D:\毕设2\启动系统.bat"),
        ("5", "访问系统", "http://127.0.0.1:5173/#/login；接口文档 http://127.0.0.1:8000/docs"),
    ])
    doc.add_heading("八、访问地址与测试账号", level=1)
    table(doc, ["项目", "内容"], [
        ("项目访问地址", "http://127.0.0.1:5173/#/login"),
        ("后台接口地址", "http://127.0.0.1:8000/docs"),
        ("测试账号", "待填写；文档不写浏览器自动填充或非确认账号"),
    ])
    save_doc(doc, OUT_ACHIEVEMENTS / f"{PROJECT}-项目部署文档.docx")


def build_report_book() -> None:
    doc = Document()
    title(doc, f"{PROJECT} 项目报告书")
    doc.add_heading("摘要", level=1)
    para(doc, f"{PROJECT} 是一个面向职业技能学习场景的本地化 RAG 知识培训系统。系统采用 React、TypeScript、Vite 构建前端学习端和管理端，后端采用 FastAPI 提供 REST API，使用 SQLite 保存用户、课程、题库、社区和反馈等业务数据，使用 ChromaDB 存储知识库向量检索数据，并接入 DeepSeek 大模型实现智能问答与实操指导生成。系统当前已实现岗位课程选择、知识库学习、智能问答、实操指导、模拟练习、错题本、个人中心、社区问答、意见反馈以及后台用户、岗位、知识库、题库、反馈和成绩管理等功能。项目仍处于部分完成状态，存在局部中文乱码、启动速度偏慢、外部模型和邮箱服务依赖、自动化测试不足等问题。本文档基于当前源码和本地运行数据如实记录设计目标、技术方案、功能实现、测试情况和后续优化方向，不包含伪造运营数据。")
    para(doc, "关键词：RAG；职业培训；FastAPI；React；知识库；智能问答")
    doc.add_heading("第一章 绪论", level=1)
    doc.add_heading("1.1 项目背景", level=2)
    para(doc, "职业技能学习通常涉及大量岗位知识、实操步骤和题库练习。传统学习平台如果只提供静态资料，学员在遇到具体问题时仍需要自行检索和归纳。RAG 技术能够把本地知识库检索结果与大模型生成能力结合起来，使系统在回答问题时尽量依据课程资料，降低纯大模型回答偏离资料的风险。")
    doc.add_heading("1.2 项目目标", level=2)
    bullets(doc, [
        "建立岗位、课程、知识库和题库之间的管理关系。",
        "为学员提供学习、问答、实操指导、练习、错题和反馈闭环。",
        "为教师或管理员提供用户、资料、社区、反馈和成绩管理能力。",
        "在文档和页面中坚持真实数据原则，无成绩记录时不展示伪造趋势。",
    ])
    doc.add_heading("第二章 需求分析", level=1)
    para(doc, "系统主要面向两类用户：学员和管理员。学员关注如何按职业方向学习、提问、练习和查看个人结果；管理员关注如何维护职业岗位、上传资料、处理反馈、管理社区和账号状态。")
    table(doc, ["用户角色", "核心需求"], [
        ("学员", "选择职业和课程，学习知识点，使用智能问答，生成实操指导，完成练习，查看错题和成绩变化，参与社区问答，提交反馈。"),
        ("管理员/教师", "维护岗位与课程，上传知识库和题库，管理用户状态，处理反馈和社区内容，查看真实练习成绩。"),
    ])
    doc.add_heading("第三章 系统设计", level=1)
    para(doc, "系统采用前后端分离结构。前端通过 Hash Router 管理登录、岗位选择、学习空间和后台管理页面；后端按功能模块拆分路由；数据库使用 SQLite；知识库检索使用 ChromaDB；大模型调用通过 DeepSeek API 封装服务完成。")
    doc.add_heading("3.1 数据设计", level=2)
    table(doc, ["数据对象", "说明"], [
        ("User", "保存用户账号、邮箱、角色、启用状态。"),
        ("Position/Course", "保存岗位和课程信息。"),
        ("KnowledgePoint/CourseMaterial/DocChunk", "保存知识点、上传资料和文档切片。"),
        ("Question/WrongQuestion", "保存题库题目和错题记录。"),
        ("CommunityQuestion/CommunityAnswer", "保存社区问题与回答。"),
        ("Feedback", "保存用户反馈和处理状态。"),
        ("QASession/QAMessage", "保存智能问答会话与消息。"),
    ])
    doc.add_heading("第四章 功能实现", level=1)
    for item in MODULES:
        doc.add_heading(item[0], level=2)
        para(doc, f"{item[1]}当前完成情况为：{item[2]}。{item[3]}")
    doc.add_heading("第五章 测试与运行情况", level=1)
    para(doc, "本轮文档生成前执行了前端 lint、前端 build 和后端 compileall，均通过。后端 pytest 因当前 Python 环境缺少 pytest 模块未执行。生成文档时前后端端口未响应，因此端口连通记录为未通过/服务未启动。")
    table(doc, ["验证项", "结果"], [
        ("npm.cmd run lint", "通过"),
        ("npm.cmd run build", "通过"),
        (r"python -m compileall -q backend\app", "通过"),
        ("python -m pytest", "未执行，缺少 pytest 模块"),
        ("5173/8000 端口连通", "未通过，生成文档时无法连接"),
    ])
    doc.add_heading("第六章 总结与展望", level=1)
    para(doc, "本项目完成了一个职业技能培训系统的主要工程框架和核心功能闭环，能够体现前后端分离、RAG 检索、大模型调用、题库练习和后台管理等综合实践内容。后续需要继续修复局部乱码，补齐端到端测试，优化启动速度，并增强知识库解析与异常降级能力。")
    doc.add_heading("参考文献", level=1)
    refs = [
        "FastAPI Documentation. https://fastapi.tiangolo.com/",
        "React Documentation. https://react.dev/",
        "Vite Documentation. https://vite.dev/",
        "SQLAlchemy Documentation. https://www.sqlalchemy.org/",
        "Chroma Documentation. https://docs.trychroma.com/",
        "PyMuPDF Documentation. https://pymupdf.readthedocs.io/",
        "DeepSeek API Documentation. https://api-docs.deepseek.com/",
        "TypeScript Documentation. https://www.typescriptlang.org/docs/",
        "Tailwind CSS Documentation. https://tailwindcss.com/docs",
        "SQLite Documentation. https://www.sqlite.org/docs.html",
    ]
    for index, ref in enumerate(refs, 1):
        doc.add_paragraph(f"[{index}] {ref}")
    doc.add_heading("附录", level=1)
    bullets(doc, [
        r"项目源码目录：D:\毕设2\frontend 与 D:\毕设2\backend。",
        r"本地数据库：D:\毕设2\backend\training.db.runtime。",
        r"启动脚本：D:\毕设2\启动系统.bat。",
        "本文档不包含真实密钥、邮箱授权码或 DeepSeek API Key。",
    ])
    save_doc(doc, OUT_REPORTS / f"{PROJECT}-项目报告书.docx")


def verify_outputs(paths: list[Path]) -> None:
    secret_patterns = [
        re.compile(r"sk-[A-Za-z0-9_-]{12,}"),
        re.compile(r"Bearer\\s+[A-Za-z0-9._-]{20,}", re.I),
    ]
    for path in paths:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = "\n".join(parts)
        if PROJECT not in text and "项目测试报告" not in text and "项目部署文档" not in text:
            raise RuntimeError(f"项目名或标题校验失败：{path}")
        if "待填写" not in text and path.name not in (f"{PROJECT}-PRD文档.docx", f"{PROJECT}-项目报告书.docx"):
            raise RuntimeError(f"待填写字段校验失败：{path}")
        for pattern in secret_patterns:
            if pattern.search(text):
                raise RuntimeError(f"疑似敏感密钥写入文档：{path}")


def main() -> None:
    if not TARGET.exists():
        raise SystemExit(f"目标目录不存在：{TARGET}")
    shutil.copytree(TARGET, BACKUP)
    for old in TARGET.rglob("智慧物流系统-*"):
        if old.is_file():
            old.chmod(stat.S_IWRITE)
            old.unlink()
    OUT_ACHIEVEMENTS.mkdir(parents=True, exist_ok=True)
    OUT_REPORTS.mkdir(parents=True, exist_ok=True)

    stats = get_stats()
    build_prd()
    build_architecture()
    build_completion()
    build_test_report(stats)
    build_deploy_doc()
    build_report_book()

    outputs = sorted(list(OUT_ACHIEVEMENTS.glob(f"{PROJECT}-*.docx")) + list(OUT_REPORTS.glob(f"{PROJECT}-*.docx")))
    verify_outputs(outputs)
    print(f"BACKUP={BACKUP}")
    for path in outputs:
        print(f"OUTPUT={path}")


if __name__ == "__main__":
    main()
